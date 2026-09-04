from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(__file__).with_name("bao-cao-tien-do-du-an.docx")


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    props.append(node)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(document, text, level=1):
    heading = document.add_heading(text, level=level)
    heading.style.font.name = "Arial"
    heading.style.font.color.rgb = RGBColor(31, 78, 121)
    return heading


def add_bullets(document, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = document.add_paragraph(style=style)
        p.add_run(item)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.7)
section.bottom_margin = Cm(1.7)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("BÁO CÁO TỔNG HỢP TIẾN ĐỘ DỰ ÁN")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(31, 78, 121)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Tổng hợp từ danh sách công việc được cung cấp").italic = True

add_heading(doc, "1. Thống kê đầu mục dự án")
doc.add_paragraph("Tổng số: 10 đầu việc — Đã làm: 8 (80%) — Chưa làm: 2 (20%).")

summary = [
    ("Nhóm", "Số đầu việc", "Đã làm", "Chưa làm"),
    ("Admin", 2, 2, 0),
    ("Chức năng", 4, 2, 2),
    ("Ổn định & hiệu năng", 1, 1, 0),
    ("Giao diện", 1, 1, 0),
    ("Thanh toán", 2, 2, 0),
    ("Tổng", 10, 8, 2),
]
table = doc.add_table(rows=len(summary), cols=4)
table.style = "Table Grid"
for i, row in enumerate(summary):
    for j, value in enumerate(row):
        set_cell_text(table.cell(i, j), value, bold=(i in (0, len(summary) - 1)), color=(255, 255, 255) if i == 0 else None)
        if i == 0:
            shade(table.cell(i, j), "1F4E79")
        elif i == len(summary) - 1:
            shade(table.cell(i, j), "D9EAF7")

add_heading(doc, "2. Người phụ trách dự án")
add_bullets(doc, [
    "Khang: tham gia cả 10 đầu việc; phụ trách chính 9 đầu việc và phối hợp 1 đầu việc.",
    "Hùng: phối hợp với Khang trong đầu việc tối ưu Chrome.",
])

add_heading(doc, "3. Chi tiết dự án và deadline")
detail = [
    ("STT", "Nhóm", "Đầu việc", "Phụ trách", "Bắt đầu", "Deadline", "Trạng thái"),
    (1, "Admin", "Quản lý người dùng, cấp mã gia hạn doanh nghiệp", "Khang", "27/08", "30/08", "Đã làm"),
    (2, "Admin", "Phân quyền", "Khang", "27/08", "30/08", "Đã làm"),
    (3, "Chức năng", "Voice call qua Messenger", "Khang", "19/08", "19/08", "Đã làm"),
    (4, "Chức năng", "Call video qua fanpage", "Khang", "Chưa có", "Chưa có", "Chưa làm"),
    (5, "Chức năng", "Call video qua Messenger", "Khang", "24/08", "29/08", "Đã làm"),
    (6, "Chức năng", "Đặt biệt danh cho khách hàng", "Khang", "01/09", "05/09", "Chưa làm"),
    (7, "Ổn định & hiệu năng", "Tối ưu Chrome", "Hùng - Khang", "30/08", "31/08", "Đã làm"),
    (8, "Giao diện", "Giao diện admin quản lý và gia hạn doanh nghiệp", "Khang", "24/08", "30/08", "Đã làm"),
    (9, "Thanh toán", "Chỉnh sửa giá gói", "Khang", "22/08", "24/08", "Đã làm"),
    (10, "Thanh toán", "Quét mã QR thanh toán", "Khang", "22/08", "24/08", "Đã làm"),
]
table = doc.add_table(rows=len(detail), cols=7)
table.style = "Table Grid"
widths = [0.7, 2.2, 6.4, 2.0, 1.6, 1.6, 1.8]
for i, row in enumerate(detail):
    for j, value in enumerate(row):
        set_cell_text(table.cell(i, j), value, bold=(i == 0), color=(255, 255, 255) if i == 0 else None, size=8)
        table.cell(i, j).width = Cm(widths[j])
        if i == 0:
            shade(table.cell(i, j), "1F4E79")
        elif row[-1] == "Chưa làm":
            shade(table.cell(i, j), "FFF2CC")

p = doc.add_paragraph()
r = p.add_run("Lưu ý: ")
r.bold = True
p.add_run("Dữ liệu không có năm và ngày cập nhật báo cáo nên chưa thể kết luận đầu việc ‘Đặt biệt danh’ đã trễ hạn. ‘Call video qua fanpage’ chưa có ngày bắt đầu và deadline, cần bổ sung kế hoạch.")

add_heading(doc, "4. Nếu dự án bị chậm tiến độ thì xử lý như thế nào?")
add_bullets(doc, [
    "Xác định đầu việc đang chậm và nguyên nhân cụ thể.",
    "Người phụ trách báo ngay cho quản lý: phần đã làm, phần còn lại và khó khăn.",
    "Chia nhỏ việc còn lại, ưu tiên phần quan trọng hoặc ảnh hưởng đến khách hàng.",
    "Bổ sung người hỗ trợ hoặc điều chỉnh phạm vi công việc nếu cần.",
    "Thống nhất deadline mới và cập nhật tiến độ hằng ngày đến khi hoàn thành.",
    "Thông báo sớm cho các bên liên quan nếu ảnh hưởng đến kế hoạch chung.",
], numbered=True)

add_heading(doc, "5. Các nguyên nhân có thể dẫn đến chậm tiến độ")
add_bullets(doc, [
    "Yêu cầu chưa rõ hoặc thay đổi trong lúc thực hiện.",
    "Chưa đặt ngày bắt đầu và deadline cho đầu việc.",
    "Một người phụ trách quá nhiều đầu việc cùng lúc.",
    "Phát sinh lỗi kỹ thuật ở chức năng gọi điện/video hoặc kết nối Facebook/Messenger.",
    "Thiếu người phối hợp, thiếu dữ liệu kiểm thử hoặc phải chờ bên thứ ba.",
    "Ước lượng thời gian chưa sát khối lượng thực tế.",
    "Không cập nhật và cảnh báo tiến độ sớm.",
])

add_heading(doc, "6. Đánh giá và đề xuất chung")
doc.add_paragraph("Tiến độ tổng thể đạt 80%. Các nhóm Admin, Ổn định & hiệu năng, Giao diện và Thanh toán đã hoàn thành toàn bộ đầu việc trong danh sách.")
add_bullets(doc, [
    "Bổ sung người hỗ trợ, ngày bắt đầu và deadline cho Call video qua fanpage.",
    "Xác nhận tình trạng thực tế và kế hoạch hoàn thành chức năng Đặt biệt danh.",
    "Cập nhật bảng tiến độ hằng ngày hoặc hằng tuần; thêm các cột Mức độ ưu tiên, % hoàn thành và Lý do chậm.",
], numbered=True)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Báo cáo tiến độ dự án").font.size = Pt(8)

doc.save(OUT)
print(OUT)
